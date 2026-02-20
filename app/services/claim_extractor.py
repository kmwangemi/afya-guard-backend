"""
SHA Claim File Extraction Service
Processes uploaded SHA claim forms (PDF, Excel, DOCX, CSV) and extracts structured data
aligned to the Social Health Insurance Act, 2023 claim form layout.
"""

import hashlib
import re
from datetime import datetime
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import pdfplumber
from docx import Document

# ---------------------------------------------------------------------------
# Data model for a fully-parsed SHA claim
# ---------------------------------------------------------------------------


class SHAClaimData:
    """
    Structured container that mirrors every section of the SHA claim form.
    All fields default to None so callers can distinguish "not found" from
    "found but empty".
    """

    def __init__(self):
        # ── PART I: Health Care Provider ────────────────────────────────────
        self.claim_number: Optional[str] = None
        self.provider_id: Optional[str] = None  # Field 1
        self.provider_name: Optional[str] = None  # Field 2

        # ── PART II: Patient Details ─────────────────────────────────────────
        self.patient_last_name: Optional[str] = None  # Field 3 (Last)
        self.patient_first_name: Optional[str] = None  # Field 3 (First)
        self.patient_middle_name: Optional[str] = None  # Field 3 (Middle)
        self.sha_number: Optional[str] = (
            None  # Field 4 – Social Health Authority Number
        )
        self.residence: Optional[str] = None  # Field 5
        self.other_insurance: Optional[str] = None  # Field 6 – other health insurance
        self.relationship_to_principal: Optional[str] = None  # Field 7

        # ── PART III: Patient Visit Details ──────────────────────────────────
        # Referral – Field 8
        self.was_referred: Optional[bool] = None
        self.referral_provider: Optional[str] = None

        # Visit metadata
        self.visit_type: Optional[str] = None  # Inpatient | Outpatient | Day-care
        self.visit_admission_date: Optional[datetime] = None
        self.op_ip_number: Optional[str] = None  # OP/IP Number
        self.new_or_return_visit: Optional[str] = None  # New | Return
        self.discharge_date: Optional[datetime] = None
        self.rendering_physician: Optional[str] = None  # Name + Registration No

        # Accommodation – Field 8 (continued)
        # Valid values: Female Medical, Male Medical, Female Surgical, Male Surgical,
        # Gynaecology, Maternity, NBU, Psychiatric Unit, Burns, ICU, HDU, NICU, Isolation
        self.accommodation_type: Optional[str] = None

        # Patient Disposition – Field 9
        # Valid values: Improved | Recovered | LAMA | Absconded | Died
        self.patient_disposition: Optional[str] = None

        # Referral on discharge – Field 10
        self.discharge_referral_institution: Optional[str] = None
        self.discharge_referral_reason: Optional[str] = None

        # Diagnoses – Fields 11 & 12
        self.admission_diagnosis: Optional[str] = None
        self.discharge_diagnosis: Optional[str] = None
        self.icd11_code: Optional[str] = None
        self.related_procedure: Optional[str] = None
        self.procedure_date: Optional[datetime] = None

        # ── PART IV: SHA Health Benefits Table (Field 14) ────────────────────
        # Each item is a dict with keys:
        #   admission_date, discharge_date, case_code, icd11_procedure_code,
        #   description, preauth_no, bill_amount, claim_amount
        self.benefit_lines: List[Dict[str, Any]] = []

        # Totals derived from benefit lines
        self.total_bill_amount: Optional[Decimal] = None
        self.total_claim_amount: Optional[Decimal] = None

        # ── Signatures / Declarations ────────────────────────────────────────
        self.patient_authorised_name: Optional[str] = None
        self.declaration_date: Optional[datetime] = None

        # ── Extraction metadata ──────────────────────────────────────────────
        self.source_file: Optional[str] = None
        self.extraction_errors: List[str] = []
        self.extracted_at: datetime = datetime.now()

    @property
    def patient_full_name(self) -> Optional[str]:
        parts = filter(
            None,
            [
                self.patient_first_name,
                self.patient_middle_name,
                self.patient_last_name,
            ],
        )
        name = " ".join(parts).strip()
        return name if name else None

    def to_dict(self) -> Dict[str, Any]:
        return {
            # Provider
            "claim_number": self.claim_number,
            "provider_id": self.provider_id,
            "provider_name": self.provider_name,
            # Patient
            "patient_last_name": self.patient_last_name,
            "patient_first_name": self.patient_first_name,
            "patient_middle_name": self.patient_middle_name,
            "patient_full_name": self.patient_full_name,
            "sha_number": self.sha_number,
            "residence": self.residence,
            "other_insurance": self.other_insurance,
            "relationship_to_principal": self.relationship_to_principal,
            # Visit
            "was_referred": self.was_referred,
            "referral_provider": self.referral_provider,
            "visit_type": self.visit_type,
            "visit_admission_date": (
                self.visit_admission_date.isoformat()
                if self.visit_admission_date
                else None
            ),
            "op_ip_number": self.op_ip_number,
            "new_or_return_visit": self.new_or_return_visit,
            "discharge_date": (
                self.discharge_date.isoformat() if self.discharge_date else None
            ),
            "rendering_physician": self.rendering_physician,
            "accommodation_type": self.accommodation_type,
            "patient_disposition": self.patient_disposition,
            # Discharge referral
            "discharge_referral_institution": self.discharge_referral_institution,
            "discharge_referral_reason": self.discharge_referral_reason,
            # Diagnoses
            "admission_diagnosis": self.admission_diagnosis,
            "discharge_diagnosis": self.discharge_diagnosis,
            "icd11_code": self.icd11_code,
            "related_procedure": self.related_procedure,
            "procedure_date": (
                self.procedure_date.isoformat() if self.procedure_date else None
            ),
            # Benefits
            "benefit_lines": self.benefit_lines,
            "total_bill_amount": (
                str(self.total_bill_amount)
                if self.total_bill_amount is not None
                else None
            ),
            "total_claim_amount": (
                str(self.total_claim_amount)
                if self.total_claim_amount is not None
                else None
            ),
            # Declarations
            "patient_authorised_name": self.patient_authorised_name,
            "declaration_date": (
                self.declaration_date.isoformat() if self.declaration_date else None
            ),
            # Metadata
            "source_file": self.source_file,
            "extraction_errors": self.extraction_errors,
            "extracted_at": self.extracted_at.isoformat(),
        }


# ---------------------------------------------------------------------------
# Regex patterns aligned to the SHA form's actual printed labels
# ---------------------------------------------------------------------------

FIELD_PATTERNS: Dict[str, str] = {
    # Header
    "claim_number": r"CLAIM\s*NO[:\s]+([A-Z0-9\-/]+)",
    # Part I – Provider
    "provider_id": r"Health\s*Provider\s*Identification\s*Number[:\s]+([A-Z0-9\-]+)",
    "provider_name": r"Name\s*of\s*Health\s*Care\s*Provider[/\s]*Facility[:\s]+([^\n]+)",
    # Part II – Patient name (form splits into three separate lines)
    "patient_last_name": r"Last\s*Name[:\s]+([^\n]+)",
    "patient_first_name": r"First\s*Name[:\s]+([^\n]+)",
    "patient_middle_name": r"Middle\s*Name[:\s]+([^\n]+)",
    # Part II – Other patient details
    "sha_number": r"Social\s*Health\s*Authority\s*Number[:\s]+([A-Z0-9\-]+)",
    "residence": r"Residence[:\s]+([^\n]+)",
    "other_insurance": r"Do\s*you\s*have\s*another\s*Health\s*Insurance[^:]*:[:\s]+([^\n]+)",
    "relationship_to_principal": r"Relationship\s*to\s*the\s*Principal[:\s]+([^\n]+)",
    # Part III – Referral
    "referral_provider": r"Name\s*of\s*referring\s*Health\s*Care\s*Provider[/\s]*Facility[:\s]+([^\n]+)",
    # Part III – Visit info
    "visit_admission_date": r"Visit[/\s]*Admission\s*Date[:\s]+([0-9]{1,2}[/\-\.][0-9]{1,2}[/\-\.][0-9]{2,4})",
    "op_ip_number": r"OP[/\s]*IP\s*No\.?[:\s]+([A-Z0-9\-/]+)",
    "new_or_return_visit": r"New[/\s]*Return\s*Visit[:\s]+(\w+)",
    "discharge_date": r"Discharge\s*Date[:\s]+([0-9]{1,2}[/\-\.][0-9]{1,2}[/\-\.][0-9]{2,4})",
    "rendering_physician": r"Rendering\s*Physician\s*Name\s*and\s*Registration\s*No[:\s]+([^\n]+)",
    "accommodation_type": r"Type\s*of\s*Accommodation[:\s]+([^\n]+)",
    # Field 10 – Discharge referral
    "discharge_referral_institution": r"Name\s*of\s*Referral\s*Institution[:\s]+([^\n]+)",
    "discharge_referral_reason": r"Reason[/\s]*s?\s*for\s*referral[:\s]+([^\n]+)",
    # Fields 11 & 12 – Diagnoses
    "admission_diagnosis": r"Admission\s*Diagnosis[/es]*[:\s]+([^\n]+)",
    "discharge_diagnosis": r"Discharge\s*Diagnosis[/es]*[:\s]+([^\n]+)",
    "icd11_code": r"ICD[- ]?11\s*Code[/s]*[:\s]+([A-Z][A-Z0-9\.\-]+)",
    "related_procedure": r"Related\s*Procedure[/s]*\s*(?:\(if\s*any\))?[:\s]+([^\n]+)",
    "procedure_date": r"Date\s*of\s*Procedure[:\s]+([0-9]{1,2}[/\-\.][0-9]{1,2}[/\-\.][0-9]{2,4})",
    # Declaration
    "patient_authorised_name": r"Names?\s*\(Majina\)[:\s]+([^\n]+)",
    "declaration_date": r"Date\s*\(Tarehe\)[:\s]+([0-9]{1,2}[/\-\.][0-9]{1,2}[/\-\.][0-9]{2,4})",
}

# SHA benefit table column header variants (lowercase)
SHA_BENEFIT_COLUMN_MAP: Dict[str, List[str]] = {
    "admission_date": ["date of admission", "admission date", "admission"],
    "discharge_date": ["date of discharge", "discharge date", "discharge"],
    "case_code": ["case code", "case"],
    "icd11_procedure_code": ["icd 11", "icd11", "procedure code", "icd 11/procedure"],
    "description": ["description", "desc"],
    "preauth_no": ["preauth no", "preauth", "pre-auth", "pre auth", "preauthorisation"],
    "bill_amount": ["bill amount", "bill"],
    "claim_amount": ["claim amount", "claim"],
}

VALID_VISIT_TYPES = {"inpatient", "outpatient", "day-care", "day care"}
VALID_DISPOSITIONS = {
    "improved",
    "recovered",
    "lama",
    "leave against medical advice",
    "discharged against medical advice",
    "absconded",
    "died",
}
VALID_ACCOMMODATION_TYPES = {
    "female medical",
    "male medical",
    "female surgical",
    "male surgical",
    "gynaecology",
    "maternity",
    "nbu",
    "psychiatric unit",
    "burns",
    "icu",
    "hdu",
    "nicu",
    "isolation",
}


# ---------------------------------------------------------------------------
# Main extractor class
# ---------------------------------------------------------------------------


class SHAClaimExtractor:
    """
    Extract and validate SHA claim data from PDF, Excel, CSV, or DOCX files.
    All extraction is aligned to the Social Health Insurance Act 2023 claim form.
    """

    SUPPORTED_FORMATS = {".pdf", ".xlsx", ".xls", ".csv", ".docx"}

    def extract_from_file(self, file_path: str) -> SHAClaimData:
        """
        Main entry point. Returns a fully-populated SHAClaimData object.

        Args:
            file_path: Absolute or relative path to the uploaded claim file.

        Returns:
            SHAClaimData instance with all extractable fields populated.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file format is not supported.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported format '{ext}'. "
                f"Supported: {', '.join(self.SUPPORTED_FORMATS)}"
            )

        claim = SHAClaimData()
        claim.source_file = str(path.resolve())

        dispatch = {
            ".pdf": self._extract_from_pdf,
            ".xlsx": self._extract_from_excel,
            ".xls": self._extract_from_excel,
            ".csv": self._extract_from_csv,
            ".docx": self._extract_from_docx,
        }
        dispatch[ext](path, claim)

        # Compute totals from benefit lines
        self._compute_benefit_totals(claim)

        return claim

    # ── Format-specific extractors ──────────────────────────────────────────

    def _extract_from_pdf(self, path: Path, claim: SHAClaimData) -> None:
        """Extract from a scanned or digital PDF SHA claim form."""
        try:
            with pdfplumber.open(str(path)) as pdf:
                full_text = ""
                all_tables: List[List[List[Optional[str]]]] = []

                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    full_text += page_text + "\n"
                    page_tables = page.extract_tables()
                    if page_tables:
                        all_tables.extend(page_tables)

                # Text-based field extraction
                self._apply_patterns(full_text, claim)

                # Visit type is a checkbox row – detect by keyword proximity
                self._extract_visit_type_from_text(full_text, claim)

                # Referral answered YES/NO
                self._extract_referral_flag_from_text(full_text, claim)

                # Patient disposition checkboxes
                self._extract_disposition_from_text(full_text, claim)

                # Benefits table (Section 14)
                if all_tables:
                    claim.benefit_lines = self._parse_benefit_table(all_tables)

        except Exception as exc:
            claim.extraction_errors.append(f"PDF extraction error: {exc}")

    def _extract_from_excel(self, path: Path, claim: SHAClaimData) -> None:
        """
        Extract from an Excel SHA claim file.
        Handles two layouts:
          (a) Key-value rows: column A = label, column B = value
          (b) Columnar / tabular layout where headers are in row 1
        """
        try:
            xls = pd.ExcelFile(str(path))

            for sheet_name in xls.sheet_names:
                df = xls.parse(sheet_name, header=None, dtype=str)
                df = df.fillna("")

                # Try key-value layout first
                self._extract_from_kv_dataframe(df, claim)

                # Then try tabular layout (for benefit lines)
                df_header = xls.parse(sheet_name, header=0, dtype=str)
                if not df_header.empty:
                    self._extract_benefit_lines_from_dataframe(df_header, claim)

        except Exception as exc:
            claim.extraction_errors.append(f"Excel extraction error: {exc}")

    def _extract_from_csv(self, path: Path, claim: SHAClaimData) -> None:
        """Extract from a CSV representation of claim data."""
        try:
            df = pd.read_csv(str(path), dtype=str).fillna("")
            # CSV is almost always columnar
            self._extract_benefit_lines_from_dataframe(df, claim)
            # Also try KV in case it's a single-claim export
            self._extract_from_kv_dataframe(df.reset_index(drop=True), claim)
        except Exception as exc:
            claim.extraction_errors.append(f"CSV extraction error: {exc}")

    def _extract_from_docx(self, path: Path, claim: SHAClaimData) -> None:
        """Extract from a Word-document SHA claim form."""
        try:
            doc = Document(str(path))
            full_text = "\n".join(p.text for p in doc.paragraphs)
            self._apply_patterns(full_text, claim)
            self._extract_visit_type_from_text(full_text, claim)
            self._extract_referral_flag_from_text(full_text, claim)
            self._extract_disposition_from_text(full_text, claim)

            # Extract tables
            raw_tables = []
            for table in doc.tables:
                raw = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                raw_tables.append(raw)

            if raw_tables:
                claim.benefit_lines = self._parse_benefit_table(raw_tables)

                # Also try KV tables for header fields
                for raw in raw_tables:
                    for row in raw:
                        if len(row) >= 2:
                            self._map_kv_pair(row[0], row[1], claim)

        except Exception as exc:
            claim.extraction_errors.append(f"DOCX extraction error: {exc}")

    # ── Pattern-based text extraction ───────────────────────────────────────

    def _apply_patterns(self, text: str, claim: SHAClaimData) -> None:
        """Apply all FIELD_PATTERNS to the raw text and populate claim fields."""
        for field, pattern in FIELD_PATTERNS.items():
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if not match:
                continue
            raw = match.group(1).strip()
            self._set_field(claim, field, raw)

    def _set_field(self, claim: SHAClaimData, field: str, raw: str) -> None:
        """Convert raw string to the correct type and assign to claim."""
        date_fields = {
            "visit_admission_date",
            "discharge_date",
            "procedure_date",
            "declaration_date",
        }
        if field in date_fields:
            setattr(claim, field, self._parse_date(raw))
        else:
            setattr(claim, field, raw if raw else None)

    def _extract_visit_type_from_text(self, text: str, claim: SHAClaimData) -> None:
        """
        Detect visit type from checkbox markers or plain text.
        The form has: ☐ Inpatient  ☐ Outpatient  ☐ Day-care
        A ticked box may appear as ☑, ✓, [X], or just the word in context.
        """
        # Ticked checkbox patterns
        ticked = re.search(
            r"(?:☑|✓|\[X\]|✔)\s*(Inpatient|Outpatient|Day[\s\-]care)",
            text,
            re.IGNORECASE,
        )
        if ticked:
            claim.visit_type = ticked.group(1).strip().title()
            return

        # Fallback: "Visit type: Inpatient" style
        fallback = re.search(
            r"Visit\s*type[:\s]+(Inpatient|Outpatient|Day[\s\-]care)",
            text,
            re.IGNORECASE,
        )
        if fallback:
            claim.visit_type = fallback.group(1).strip().title()

    def _extract_referral_flag_from_text(self, text: str, claim: SHAClaimData) -> None:
        """
        Detect whether the patient was referred.
        The form has: ☐ NO  ☐ If, YES
        """
        ticked_no = re.search(r"(?:☑|✓|\[X\]|✔)\s*NO", text, re.IGNORECASE)
        ticked_yes = re.search(
            r"(?:☑|✓|\[X\]|✔)\s*(?:If[,\s]*)?\s*YES", text, re.IGNORECASE
        )
        if ticked_no:
            claim.was_referred = False
        elif ticked_yes:
            claim.was_referred = True

    def _extract_disposition_from_text(self, text: str, claim: SHAClaimData) -> None:
        """
        Detect patient disposition from ticked checkboxes.
        Options: Improved | Recovered | Leave Against Medical Advice | Absconded | Died
        """
        disposition_patterns = [
            ("Improved", r"(?:☑|✓|\[X\]|✔)\s*Improved"),
            ("Recovered", r"(?:☑|✓|\[X\]|✔)\s*Recovered"),
            ("Leave Against Medical Advice", r"(?:☑|✓|\[X\]|✔)\s*Leave\s*Against"),
            ("Absconded", r"(?:☑|✓|\[X\]|✔)\s*Absconded"),
            ("Died", r"(?:☑|✓|\[X\]|✔)\s*Died"),
        ]
        for label, pattern in disposition_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                claim.patient_disposition = label
                return

    # ── Benefit table extraction ─────────────────────────────────────────────

    def _parse_benefit_table(
        self, tables: List[List[List[Optional[str]]]]
    ) -> List[Dict[str, Any]]:
        """
        Parse Section 14 (SHA Health Benefits) table.
        Columns: Date of Admission | Date of Discharge | Case Code |
                 ICD 11/Procedure Code | Description | Preauth No. |
                 Bill Amount | Claim Amount
        """
        lines: List[Dict[str, Any]] = []

        for table in tables:
            if not table or len(table) < 2:
                continue

            header_row = [str(c).lower().strip() if c else "" for c in table[0]]

            # Build column-index map
            col_map: Dict[str, int] = {}
            for our_col, variants in SHA_BENEFIT_COLUMN_MAP.items():
                for i, header_cell in enumerate(header_row):
                    if any(v in header_cell for v in variants):
                        col_map[our_col] = i
                        break

            # This table must have at least one amount column to be the benefits table
            if "claim_amount" not in col_map and "bill_amount" not in col_map:
                continue

            for row in table[1:]:
                # Skip completely empty rows and the "Total" summary row
                if not row:
                    continue
                cells = [str(c).strip() if c else "" for c in row]
                if all(c == "" for c in cells):
                    continue
                if any("total" in c.lower() for c in cells):
                    continue

                line: Dict[str, Any] = {}
                for field, idx in col_map.items():
                    if idx >= len(cells):
                        continue
                    val = cells[idx]
                    if not val:
                        continue
                    if "amount" in field:
                        line[field] = self._parse_amount(val)
                    elif "date" in field:
                        line[field] = self._parse_date(val)
                    else:
                        line[field] = val

                if any(v for v in line.values() if v is not None):
                    lines.append(line)

        return lines

    def _compute_benefit_totals(self, claim: SHAClaimData) -> None:
        """Sum bill_amount and claim_amount across all benefit lines."""
        total_bill = Decimal("0")
        total_claim = Decimal("0")
        has_bill = has_claim = False

        for line in claim.benefit_lines:
            if isinstance(line.get("bill_amount"), Decimal):
                total_bill += line["bill_amount"]
                has_bill = True
            if isinstance(line.get("claim_amount"), Decimal):
                total_claim += line["claim_amount"]
                has_claim = True

        if has_bill:
            claim.total_bill_amount = total_bill
        if has_claim:
            claim.total_claim_amount = total_claim

    # ── Key-value / dataframe extraction ────────────────────────────────────

    def _extract_from_kv_dataframe(self, df: pd.DataFrame, claim: SHAClaimData) -> None:
        """
        Handle spreadsheets where column 0 = label, column 1 = value.
        """
        if df.shape[1] < 2:
            return
        for _, row in df.iterrows():
            label = str(row.iloc[0]).strip()
            value = str(row.iloc[1]).strip()
            if label and value:
                self._map_kv_pair(label, value, claim)

    def _map_kv_pair(self, label: str, value: str, claim: SHAClaimData) -> None:
        """Map a single label/value pair to the correct claim field."""
        if not label or not value or value.lower() in {"nan", "none", ""}:
            return

        ll = label.lower()

        # Provider fields
        if "provider" in ll and "id" in ll:
            claim.provider_id = claim.provider_id or value
        elif "provider" in ll and ("name" in ll or "facility" in ll):
            claim.provider_name = claim.provider_name or value

        # Patient name fields
        elif "last" in ll and "name" in ll:
            claim.patient_last_name = claim.patient_last_name or value
        elif "first" in ll and "name" in ll:
            claim.patient_first_name = claim.patient_first_name or value
        elif "middle" in ll and "name" in ll:
            claim.patient_middle_name = claim.patient_middle_name or value
        elif "patient" in ll and "name" in ll:
            # Full name in one cell – try to split on comma or space
            claim.patient_last_name = claim.patient_last_name or value

        # SHA / membership number
        elif "sha" in ll and ("number" in ll or "no" in ll):
            claim.sha_number = claim.sha_number or value

        # Dates
        elif "admission" in ll or ("visit" in ll and "date" in ll):
            claim.visit_admission_date = claim.visit_admission_date or self._parse_date(
                value
            )
        elif "discharge" in ll and "date" in ll:
            claim.discharge_date = claim.discharge_date or self._parse_date(value)
        elif "procedure" in ll and "date" in ll:
            claim.procedure_date = claim.procedure_date or self._parse_date(value)

        # Diagnosis
        elif "admission" in ll and "diagnosis" in ll:
            claim.admission_diagnosis = claim.admission_diagnosis or value
        elif "discharge" in ll and "diagnosis" in ll:
            claim.discharge_diagnosis = claim.discharge_diagnosis or value
        elif "icd" in ll:
            claim.icd11_code = claim.icd11_code or value

        # Visit type
        elif "visit" in ll and "type" in ll:
            claim.visit_type = claim.visit_type or value

        # Accommodation
        elif "accommodation" in ll:
            claim.accommodation_type = claim.accommodation_type or value

        # Disposition
        elif "disposition" in ll:
            claim.patient_disposition = claim.patient_disposition or value

        # Physician
        elif "physician" in ll or ("rendering" in ll and "doctor" in ll):
            claim.rendering_physician = claim.rendering_physician or value

        # Preauth
        elif "preauth" in ll or "pre-auth" in ll:
            # May be a single preauth for non-table forms
            if claim.benefit_lines:
                claim.benefit_lines[0].setdefault("preauth_no", value)

        # Referral
        elif "referral" in ll and "institution" in ll:
            claim.discharge_referral_institution = (
                claim.discharge_referral_institution or value
            )
        elif "referral" in ll and "reason" in ll:
            claim.discharge_referral_reason = claim.discharge_referral_reason or value
        elif "referring" in ll:
            claim.referral_provider = claim.referral_provider or value

        # Claim / OP-IP number
        elif "claim" in ll and "no" in ll:
            claim.claim_number = claim.claim_number or value
        elif "op" in ll and "ip" in ll:
            claim.op_ip_number = claim.op_ip_number or value

        # Residence
        elif "residence" in ll:
            claim.residence = claim.residence or value

        # Other insurance
        elif "insurance" in ll and ("other" in ll or "another" in ll):
            claim.other_insurance = claim.other_insurance or value

        # Relationship to principal
        elif "relationship" in ll and "principal" in ll:
            claim.relationship_to_principal = claim.relationship_to_principal or value

    def _extract_benefit_lines_from_dataframe(
        self, df: pd.DataFrame, claim: SHAClaimData
    ) -> None:
        """
        Extract benefit lines when the spreadsheet is in columnar format
        (one benefit line per row, headers in row 0).
        """
        if df.empty:
            return

        headers = [str(c).lower().strip() for c in df.columns]

        # Build column-index map
        col_map: Dict[str, str] = {}
        for our_col, variants in SHA_BENEFIT_COLUMN_MAP.items():
            for col in df.columns:
                if any(v in str(col).lower() for v in variants):
                    col_map[our_col] = col
                    break

        if "claim_amount" not in col_map and "bill_amount" not in col_map:
            return  # Not the benefits table

        for _, row in df.iterrows():
            line: Dict[str, Any] = {}
            for field, col in col_map.items():
                val = str(row[col]).strip() if pd.notna(row[col]) else ""
                if not val or val.lower() in {"nan", "none"}:
                    continue
                if "amount" in field:
                    line[field] = self._parse_amount(val)
                elif "date" in field:
                    line[field] = self._parse_date(val)
                else:
                    line[field] = val

            if any(v for v in line.values() if v is not None):
                claim.benefit_lines.append(line)

    # ── Type converters ──────────────────────────────────────────────────────

    @staticmethod
    def _parse_date(value: Any) -> Optional[datetime]:
        """Parse a date from string or datetime objects."""
        if value is None:
            return None
        if isinstance(value, datetime):
            return value
        if isinstance(value, pd.Timestamp):
            return value.to_pydatetime()

        value_str = str(value).strip()
        if not value_str or value_str.lower() in {"nan", "none", "n/a"}:
            return None

        formats = [
            "%d/%m/%Y",
            "%d-%m-%Y",
            "%d.%m.%Y",
            "%Y-%m-%d",
            "%Y/%m/%d",
            "%m/%d/%Y",
            "%d %b %Y",
            "%d %B %Y",
            "%d/%m/%y",
            "%d-%m-%y",
        ]
        for fmt in formats:
            try:
                return datetime.strptime(value_str, fmt)
            except ValueError:
                continue
        return None

    @staticmethod
    def _parse_amount(value: Any) -> Optional[Decimal]:
        """Parse a monetary amount, stripping KSh symbols and commas."""
        if value is None:
            return None
        raw = str(value).replace("KSh", "").replace("Ksh", "").replace("KES", "")
        raw = raw.replace(",", "").strip()
        # Remove any trailing non-numeric chars except dot
        raw = re.sub(r"[^\d\.]", "", raw)
        if not raw:
            return None
        try:
            return Decimal(raw)
        except InvalidOperation:
            return None

    # ── Validation ───────────────────────────────────────────────────────────

    def validate(self, claim: SHAClaimData) -> Tuple[bool, List[str]]:
        """
        Validate an extracted SHAClaimData object.

        Returns:
            (is_valid, list_of_error_messages)
        """
        errors: List[str] = []

        # ── Required fields ──────────────────────────────────────────────────
        required: Dict[str, str] = {
            "provider_id": "Provider Identification Number (Part I, Field 1)",
            "provider_name": "Name of Health Care Provider (Part I, Field 2)",
            "sha_number": "Social Health Authority Number (Part II, Field 3)",
            "visit_admission_date": "Visit/Admission Date (Part III)",
            "visit_type": "Visit Type – Inpatient/Outpatient/Day-care (Part III)",
        }
        for attr, description in required.items():
            if not getattr(claim, attr):
                errors.append(f"Missing required field: {description}")

        # ── Visit type enum ──────────────────────────────────────────────────
        if claim.visit_type and claim.visit_type.lower() not in VALID_VISIT_TYPES:
            errors.append(
                f"Invalid visit_type '{claim.visit_type}'. "
                f"Must be one of: {', '.join(VALID_VISIT_TYPES)}"
            )

        # ── Accommodation type enum ──────────────────────────────────────────
        if (
            claim.accommodation_type
            and claim.accommodation_type.lower() not in VALID_ACCOMMODATION_TYPES
        ):
            errors.append(
                f"Unrecognised accommodation_type '{claim.accommodation_type}'. "
                f"Valid types: {', '.join(sorted(VALID_ACCOMMODATION_TYPES))}"
            )

        # ── Patient disposition enum ─────────────────────────────────────────
        if (
            claim.patient_disposition
            and claim.patient_disposition.lower() not in VALID_DISPOSITIONS
        ):
            errors.append(
                f"Unrecognised patient_disposition '{claim.patient_disposition}'"
            )

        # ── Date logic ───────────────────────────────────────────────────────
        if claim.visit_admission_date and claim.discharge_date:
            if claim.discharge_date < claim.visit_admission_date:
                errors.append(
                    "discharge_date cannot be before visit_admission_date "
                    "(potential fraud signal)"
                )
            # Flag excessively long stays (>180 days) for review
            stay_days = (claim.discharge_date - claim.visit_admission_date).days
            if stay_days > 180:
                errors.append(
                    f"Length of stay is {stay_days} days – flagged for manual review"
                )

        if claim.procedure_date and claim.visit_admission_date:
            if claim.procedure_date < claim.visit_admission_date:
                errors.append(
                    "procedure_date is before admission date (potential fraud signal)"
                )

        # ── ICD-11 code format ───────────────────────────────────────────────
        if claim.icd11_code:
            if not re.match(
                r"^[A-Z][A-Z0-9]{1,5}(?:\.[A-Z0-9]{1,4})*$", claim.icd11_code
            ):
                errors.append(
                    f"icd11_code '{claim.icd11_code}' does not match expected ICD-11 format"
                )

        # ── SHA number format ────────────────────────────────────────────────
        if claim.sha_number:
            if not re.match(r"^[A-Z0-9\-]{5,20}$", claim.sha_number, re.IGNORECASE):
                errors.append(
                    f"sha_number '{claim.sha_number}' does not match expected format"
                )

        # ── Benefit lines ────────────────────────────────────────────────────
        if not claim.benefit_lines:
            errors.append(
                "No benefit lines found in Section 14 (SHA Health Benefits table)"
            )
        else:
            for i, line in enumerate(claim.benefit_lines, start=1):
                bill = line.get("bill_amount")
                clm = line.get("claim_amount")
                if bill is not None and clm is not None:
                    if clm > bill:
                        errors.append(
                            f"Benefit line {i}: claim_amount ({clm}) exceeds "
                            f"bill_amount ({bill}) – fraud signal"
                        )
                    if clm <= 0:
                        errors.append(f"Benefit line {i}: claim_amount must be > 0")
                if not line.get("preauth_no"):
                    errors.append(
                        f"Benefit line {i}: missing preauth_no – required for claim processing"
                    )

        # ── Claim amount vs bill amount totals ───────────────────────────────
        if claim.total_claim_amount is not None and claim.total_bill_amount is not None:
            if claim.total_claim_amount > claim.total_bill_amount:
                errors.append(
                    f"Total claim amount ({claim.total_claim_amount}) exceeds "
                    f"total bill amount ({claim.total_bill_amount}) – potential overbilling"
                )

        # ── Referral consistency ─────────────────────────────────────────────
        if claim.was_referred is True and not claim.referral_provider:
            errors.append(
                "Patient marked as referred but no referring provider recorded (Field 7)"
            )

        return len(errors) == 0, errors

    # ── Utility ──────────────────────────────────────────────────────────────

    @staticmethod
    def generate_claim_number(provider_code: str) -> str:
        """
        Generate a unique claim reference number.
        Format: CLM-<PROVIDER_CODE>-<YYYYMMDDHHMMSS>-<6-char hash>
        """
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        suffix = (
            hashlib.md5(f"{provider_code}{timestamp}".encode()).hexdigest()[:6].upper()
        )
        return f"CLM-{provider_code.upper()}-{timestamp}-{suffix}"


# ---------------------------------------------------------------------------
# Module-level singleton
# ---------------------------------------------------------------------------

sha_claim_extractor = SHAClaimExtractor()
